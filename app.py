from flask import Flask, render_template, request, redirect, url_for, session, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import pandas as pd
import os

app = Flask(__name__)
app.secret_key = "fuxing_minimalist_2025"

# --- 共用 CSS 樣式 (極簡黑白灰) ---
STYLE = '''
<style>
    body { font-family: "Microsoft JhengHei", sans-serif; background-color: #ffffff; color: #333; display: flex; justify-content: center; padding-top: 50px; }
    .container { width: 400px; border: 1px solid #ddd; padding: 40px; border-radius: 4px; background: #fff; }
    h2, h3, h4 { color: #000; text-align: center; font-weight: 600; margin-bottom: 25px; }
    input[type="text"], input[type="password"] { width: 100%; padding: 12px; margin: 10px 0; border: 1px solid #ccc; border-radius: 2px; box-sizing: border-box; }
    .btn { width: 100%; padding: 12px; background: #333; color: #fff; border: none; border-radius: 2px; cursor: pointer; font-size: 16px; margin: 10px 0; transition: 0.3s; }
    .btn:hover { background: #555; }
    .btn-outline { background: #fff; color: #333; border: 1px solid #333; }
    .btn-outline:hover { background: #f5f5f5; }
    .info-box { font-size: 14px; background: #f9f9f9; padding: 15px; border-radius: 2px; border-left: 4px solid #333; margin-bottom: 20px; }
    .radio-group { margin: 15px 0; line-height: 2.5; }
    hr { border: 0; border-top: 1px solid #eee; margin: 20px 0; }
    a { text-decoration: none; }
    .logout { display: block; text-align: center; color: #999; font-size: 13px; margin-top: 20px; }
</style>
'''

# --- 1. 登入頁面 ---
@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        pwd_tail = request.form.get('pwd_tail', '').strip()
        filename = 'students.xlsx'
        try:
            if not os.path.exists(filename): return "錯誤：找不到學生資料庫"
            db = pd.read_excel(filename, dtype=str)
            names_in_excel = db.iloc[:, 3].str.strip()
            ids_in_excel = db.iloc[:, 4].str.strip()
            match = (names_in_excel == name) & (ids_in_excel.str.endswith(str(pwd_tail)))
            user_data = db[match]
            if not user_data.empty:
                row = user_data.iloc[0]
                session['user'] = {
                    'class_name': str(row.iloc[0]), 'seat_number': str(row.iloc[1]),
                    'student_id': str(row.iloc[2]), 'name': str(row.iloc[3])
                }
                return redirect(url_for('select_group'))
            return "登入失敗：請檢查姓名與後六碼"
        except Exception as e: return f"系統錯誤：{str(e)}"

    return f'''
    {STYLE}
    <div class="container">
        <h2>復興高中選班群系統</h2>
        <form method="post">
            <input type="text" name="name" placeholder="姓名" required>
            <input type="password" name="pwd_tail" placeholder="身分證後六碼" required>
            <button type="submit" class="btn">登入</button>
        </form>
    </div>
    '''

# --- 2. 選擇班群 ---
@app.route('/select', methods=['GET', 'POST'])
def select_group():
    if 'user' not in session: return redirect(url_for('login'))
    user = session['user']
    if request.method == 'POST':
        session['selected_group'] = request.form.get('group')
        session['selection_time'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        return redirect(url_for('confirm'))
    
    return f'''
    {STYLE}
    <div class="container">
        <h3>基本資料確認</h3>
        <div class="info-box">
            {user['class_name']} 班 {user['seat_number']} 號<br>
            {user['name']} ({user['student_id']})
        </div>
        <form method="post">
            <div class="radio-group">
                <input type="radio" name="group" value="文法商班群(數A)" required> 文法商班群(數A)<br>
                <input type="radio" name="group" value="文法商班群(數B)"> 文法商班群(數B)<br>
                <input type="radio" name="group" value="理工資班群"> 理工資班群<br>
                <input type="radio" name="group" value="生醫農班群"> 生醫農班群
            </div>
            <button type="submit" class="btn">下一步</button>
        </form>
    </div>
    '''

# --- 3. 確認與操作 ---
@app.route('/confirm')
def confirm():
    if 'selected_group' not in session: return redirect(url_for('select_group'))
    return f'''
    {STYLE}
    <div class="container">
        <h3>已選擇：{session['selected_group']}</h3>
        <p style="text-align:center; font-size:14px; color:#666;">請依序完成下方步驟</p>
        <hr>
        <a href="/submit"><button class="btn btn-outline">1. 送出並存檔 (Excel)</button></a>
        <a href="/download"><button class="btn">2. 產生列印用 Word</button></a>
        <a href="/logout" class="logout">登出系統</a>
    </div>
    '''

# --- 4. 存檔邏輯 (Excel) ---
@app.route('/submit')
def submit():
    if 'user' not in session: return redirect(url_for('login'))
    user, group = session['user'], session.get('selected_group')
    file_name = 'selections.xlsx'
    data = {"班級": user['class_name'], "座號": user['seat_number'], "姓名": user['name'], "學號": user['student_id'], "選定班群": group, "時間": session.get('selection_time')}
    df = pd.DataFrame([data])
    try:
        if os.path.exists(file_name):
            old_df = pd.read_excel(file_name)
            df = pd.concat([old_df, df], ignore_index=True)
        df.to_excel(file_name, index=False)
        return f"{STYLE}<div class='container' style='text-align:center;'><p>✅ 存檔成功</p><a href='/confirm' class='btn'>返回</a></div>"
    except Exception as e: return f"存檔失敗：{e}"

# --- 5. Word 生成 ---
@app.route('/download')
def download():
    if 'user' not in session: return redirect(url_for('login'))
    user, group = session['user'], session.get('selected_group')
    try:
        doc = Document('template.docx')
        repls = {
            "{{class_name}}": str(user['class_name']), "{{seat_number}}": str(user['seat_number']),
            "{{student_id}}": str(user['student_id']), "{{name}}": str(user['name']),
            "{{selection_timestamp}}": session.get('selection_time'),
            "{{print_timestamp}}": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "{{checkbox_liberal_arts_A}}": "V" if group == "文法商班群(數A)" else " ",
            "{{checkbox_liberal_arts_B}}": "V" if group == "文法商班群(數B)" else " ",
            "{{checkbox_stem}}": "V" if group == "理工資班群" else " ",
            "{{checkbox_bio_agri}}": "V" if group == "生醫農班群" else " "
        }
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for k, v in repls.items():
                            if k in p.text:
                                p.text = p.text.replace(k, v)
                                if v == "V" and "{{checkbox_" in k:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in p.runs: run.font.size, run.font.bold = Pt(22), True
        for p in doc.paragraphs:
            for k, v in repls.items():
                if k in p.text: p.text = p.text.replace(k, v)
        out_file = f"Result_{user['student_id']}.docx"
        doc.save(out_file)
        return send_file(out_file, as_attachment=True)
    except Exception as e: return f"下載失敗：{e}"

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))
# --- 管理與統計功能 ---

@app.route('/admin')
def admin_summary():
    # 1. 讀取原始全校學生名冊
    try:
        df_students = pd.read_excel("students.xlsx")
        # 確保學號為字串格式
        df_students['學號'] = df_students['學號'].astype(str).str.strip()
    except Exception as e:
        return f"讀取 students.xlsx 失敗: {e}"

    # 2. 讀取已選擇紀錄
    if os.path.exists("selections.xlsx"):
        try:
            df_selections = pd.read_excel("selections.xlsx")
            df_selections['學號'] = df_selections['學號'].astype(str).str.strip()
            # 只取需要的欄位進行比對
            df_selections = df_selections[['學號', '選定班群', '時間']]
        except:
            df_selections = pd.DataFrame(columns=['學號', '選定班群', '時間'])
    else:
        df_selections = pd.DataFrame(columns=['學號', '選定班群', '時間'])

    # 3. 合併資料 (Left Join)
    merged = pd.merge(df_students[['班級', '座號', '學號', '姓名']], df_selections, on='學號', how='left')

    # 4. 統計人數
    total_count = len(merged)
    completed_count = df_selections['學號'].nunique()
    pending_count = total_count - completed_count

    # 5. 格式化表格內容
    merged['選定班群'] = merged['選定班群'].apply(lambda x: f'<span style="color:red; font-weight:bold;">尚未填寫</span>' if pd.isnull(x) else x)
    merged['時間'] = merged['時間'].fillna('-')

    table_html = merged.to_html(classes='table', index=False, escape=False)

    # 6. 回傳管理頁面
    return f'''
    <html>
    <head>
        <title>全校選課統計</title>
        <style>
            body {{ font-family: "Microsoft JhengHei", sans-serif; padding: 30px; background-color: #f8f9fa; }}
            .container {{ max-width: 1000px; margin: auto; background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
            .summary {{ display: flex; justify-content: space-around; background: #333; color: white; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
            .table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
            .table th, .table td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
            .table th {{ background-color: #f2f2f2; position: sticky; top: 0; }}
            .btn-download {{ display: block; width: 200px; text-align: center; background: #28a745; color: white; padding: 10px; text-decoration: none; border-radius: 5px; margin-bottom: 10px; font-weight: bold; }}
            .btn-download:hover {{ background: #218838; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h2>全校選課進度統計</h2>
            <div class="summary">
                <span>總人數：{total_count}</span>
                <span style="color: #5cb85c;">已送出：{completed_count}</span>
                <span style="color: #ff4d4d;">尚未送出：{pending_count}</span>
            </div>
            <a href="/export_excel" class="btn-download">⬇️ 下載完整 Excel 報表</a>
            <div style="overflow-y: auto; max-height: 600px;">
                {table_html}
            </div>
        </div>
    </body>
    </html>
    '''

@app.route('/export_excel')
def export_excel():
    # 讀取全校學生
    df_students = pd.read_excel("students.xlsx")
    df_students['學號'] = df_students['學號'].astype(str).str.strip()
    
    # 讀取已選擇紀錄
    if os.path.exists("selections.xlsx"):
        df_selections = pd.read_excel("selections.xlsx")
        df_selections['學號'] = df_selections['學號'].astype(str).str.strip()
        df_selections = df_selections[['學號', '選定班群', '時間']]
    else:
        df_selections = pd.DataFrame(columns=['學號', '選定班群', '時間'])

    # 合併
    report = pd.merge(df_students[['班級', '座號', '學號', '姓名']], df_selections, on='學號', how='left')
    
    # 新增狀態欄位
    report['填寫狀態'] = report['選定班群'].apply(lambda x: '已完成' if pd.notnull(x) else '尚未填寫')
    
    # 儲存
    output_path = "全校選課統計結果.xlsx"
    report.to_excel(output_path, index=False)
    
    return send_file(output_path, as_attachment=True)
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)